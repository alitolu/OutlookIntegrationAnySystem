import os
import tempfile
from typing import Dict, Optional, List
import pytesseract
from PIL import Image
import pdf2image
import cv2
import numpy as np
from docx import Document 

class AttachmentProcessor:
    def __init__(self, config=None):
        self.config = config
        
        # Temel dizin
        self.base_dir = os.path.dirname(os.path.dirname(__file__)) 
        self.temp_dir = os.path.join(self.base_dir, "temp_ocr")
        
        # temp_ocr klasörünü oluştur
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)
            
        # Desteklenen formatlar
        self.supported_formats = set(config.get("ocr", {}).get("supported_formats", [
            '.pdf', '.png', '.jpg', '.jpeg', '.tiff', '.doc', '.docx', '.xml'
        ]))
        
        # OCR ayarları
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        self.languages = '+'.join(config.get("ocr", {}).get("languages", ["tur", "eng"]))
        self.preprocess_params = config.get("ocr", {}).get("preprocess", {
            "denoise": True,
            "threshold": 128,
            "contrast": 1.5,
            "dpi": 300
        })
        self.processed_files = []  # İşlenmiş dosyaları takip etmek için liste ekle

        # Poppler path'ini ayarla
        self.poppler_path = r"C:\poppler\bin"
        if not os.path.exists(self.poppler_path):
            print(f"UYARI: Poppler bulunamadı: {self.poppler_path}")
            print("PDF işleme devre dışı kalacak")

    def process_attachment(self, attachment) -> Optional[Dict]:
        """Ekleri işle"""
        try:
            # Benzersiz dosya adı oluştur
            filename = attachment.FileName
            base_name, ext = os.path.splitext(filename)
            counter = 1
            temp_path = os.path.join(self.temp_dir, filename)
            
            # Aynı isimde dosya varsa numara ekle
            while os.path.exists(temp_path):
                new_filename = f"{base_name}_{counter}{ext}"
                temp_path = os.path.join(self.temp_dir, new_filename)
                counter += 1
                
            # Dosyayı kaydet
            attachment.SaveAsFile(temp_path)
            self.processed_files.append(temp_path)  # Dosyayı listeye ekle
            print(f"Ek kaydedildi: {temp_path}")
                
            # Dosya tipine göre işle
            ext = ext.lower()
            if ext in self.supported_formats:
                result = None
                if ext == '.pdf':
                    result = self._process_pdf(temp_path)
                elif ext in ['.tiff']:
                    result = self._process_image(temp_path)
                elif ext in ['.doc', '.docx']:
                    result = self._process_doc(temp_path)
                
                if result:
                    result_dict = {
                        'content': result,
                        'type': ext,
                        'filename': filename,
                        'temp_path': temp_path
                    }
                    self.cleanup()  # İşlem bitince temizlik yap
                    return result_dict
            
            self.cleanup()  # Hata durumunda da temizlik yap
            return None
            
        except Exception as e:
            print(f"Ek işleme hatası: {str(e)}")
            self.cleanup()  # Hata durumunda temizlik yap
            return None

    def _process_image(self, image_path: str) -> str:
        """Görüntü işleme ve OCR"""
        try:
            # Görüntüyü oku
            image = cv2.imread(image_path)
            if image is None:
                return ""
                
            # Ön işleme
            processed = self._preprocess_image(image)
            
            # OCR uygula
            text = pytesseract.image_to_string(processed, lang=self.languages)
            return text.strip()
            
        except Exception as e:
            print(f"Görüntü işleme hatası: {str(e)}")
            return ""

    def _process_pdf(self, pdf_path: str) -> str:
        """PDF işleme"""
        try:
            if not os.path.exists(self.poppler_path):
                print("PDF işleme yapılamıyor - Poppler kurulu değil")
                return ""
                
            # PDF'i görüntülere çevir - poppler path'i ekle
            pages = pdf2image.convert_from_path(
                pdf_path,
                poppler_path=self.poppler_path
            )
            results = []
            
            for i, page in enumerate(pages):
                # Her sayfa için geçici PNG dosyası oluştur
                temp_path = os.path.join(self.temp_dir, f'page_{i}.png')
                page.save(temp_path, 'PNG')
                
                # OCR uygula
                text = self._process_image(temp_path)
                if text:
                    results.append(text)
                    
                # Geçici PNG'yi sil
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                
            return "\n\n".join(results)
            
        except Exception as e:
            print(f"PDF işleme hatası: {str(e)}")
            return ""

    def _process_doc(self, file_path: str) -> str:
        """Word dokümanlarını işle"""
        try:
            # Dokümanı aç
            doc = Document(file_path)
            
            # Tüm paragrafları birleştir
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
                    
            # Tablolardaki metinleri de ekle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
                            
            return "\n\n".join(text)
            
        except Exception as e:
            print(f"Word dokümanı işleme hatası: {str(e)}")
            return ""

    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Görüntü ön işleme"""
        try:
            # Griye çevir
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Gürültü temizle
            if self.preprocess_params["denoise"]:
                gray = cv2.fastNlMeansDenoising(gray)
            
            # Kontrast ayarla    
            gray = cv2.convertScaleAbs(gray, alpha=self.preprocess_params["contrast"])
            
            # Eşikleme
            _, threshold = cv2.threshold(
                gray,
                self.preprocess_params["threshold"],
                255,
                cv2.THRESH_BINARY + cv2.THRESH_OTSU
            )
            
            return threshold
            
        except Exception as e:
            print(f"Ön işleme hatası: {str(e)}")
            return image

    def _get_extension(self, attachment) -> str:
        """Dosya uzantısını al"""
        return os.path.splitext(attachment.FileName)[1].lower()

    def cleanup(self):
        """İşlenmiş dosyaları temizle"""
        try:
            for file_path in self.processed_files:
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                        print(f"Geçici dosya silindi: {file_path}")
                    except Exception as e:
                        print(f"Dosya silinirken hata: {file_path} - {str(e)}")
            
            self.processed_files.clear()  # Listeyi temizle
            
            # PDF işlemeden kalan page_*.png dosyalarını temizle
            for file in os.listdir(self.temp_dir):
                if file.startswith('page_') and file.endswith('.png'):
                    try:
                        os.remove(os.path.join(self.temp_dir, file))
                    except:
                        pass
                        
        except Exception as e:
            print(f"Temizlik hatası: {str(e)}")

    def __del__(self):
        """Destructor - son bir temizlik yap"""
        self.cleanup()